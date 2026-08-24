"""Turn an uploaded file into plain text.

PDFs go through pypdf, DOCX files through python-docx, and plain text files are
decoded directly. Images go through Gemini vision: the embedding model is
text-only, so an image has to be transcribed before it can be indexed at all.
"""

import io

from docx import Document as DocxDocument
from google.genai import types
from pypdf import PdfReader

from app.config import GEMINI_MODEL
from app.gemini_client import get_client

PDF_MIME_TYPES = {"application/pdf"}
IMAGE_MIME_TYPES = {
    "image/png",
    "image/jpeg",
    "image/jpg",
    "image/webp",
}
DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
TEXT_MIME_TYPES = {"text/plain"}
PDF_EXTENSIONS = {".pdf"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}
DOCX_EXTENSIONS = {".docx"}
TEXT_EXTENSIONS = {".txt"}

IMAGE_PROMPT = """Extract everything a reader could learn from this image.

1. Transcribe all visible text verbatim, preserving reading order. Keep table rows on
   one line with cells separated by " | ".
2. Then describe any chart, diagram, form, screenshot or photo content, including axis
   labels, legends, values and relationships between elements.
3. Do not add commentary, guesses, or information that is not visible in the image.

If the image contains no text and no meaningful content, reply exactly: EMPTY_IMAGE"""


class ExtractionError(RuntimeError):
    """Raised when a file yields no usable text."""


def extract_pdf_text(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    except Exception as error:
        raise ExtractionError(f"Could not read the PDF: {error}") from error


def extract_docx_text(data: bytes) -> str:
    """Extract paragraph and table text from a DOCX file."""
    try:
        document = DocxDocument(io.BytesIO(data))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(part for part in parts if part).strip()
    except Exception as error:
        raise ExtractionError(f"Could not read the DOCX file: {error}") from error


def extract_plain_text(data: bytes) -> str:
    """Decode a plain text file, tolerating non-UTF-8 encodings."""
    for encoding in ("utf-8", "latin-1"):
        try:
            return data.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise ExtractionError("Could not decode this text file.")


def extract_image_text(data: bytes, mime_type: str) -> str:
    try:
        response = get_client().models.generate_content(
            model=GEMINI_MODEL,
            contents=[
                types.Part.from_bytes(data=data, mime_type=mime_type),
                IMAGE_PROMPT,
            ],
        )
    except Exception as error:
        raise ExtractionError(f"Gemini vision request failed: {error}") from error

    text = (response.text or "").strip()
    if not text or text == "EMPTY_IMAGE":
        raise ExtractionError("No readable content was found in this image.")
    return text


def detect_kind(filename: str, content_type: str | None) -> str:
    """Return "pdf", "image", "docx" or "text", or raise ValueError for anything else."""
    suffix = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    mime = (content_type or "").split(";")[0].strip().lower()

    extension_kind = (
        "pdf"
        if suffix in PDF_EXTENSIONS
        else "image"
        if suffix in IMAGE_EXTENSIONS
        else "docx"
        if suffix in DOCX_EXTENSIONS
        else "text"
        if suffix in TEXT_EXTENSIONS
        else None
    )
    mime_kind = (
        "pdf"
        if mime in PDF_MIME_TYPES
        else "image"
        if mime in IMAGE_MIME_TYPES
        else "docx"
        if mime in DOCX_MIME_TYPES
        else "text"
        if mime in TEXT_MIME_TYPES
        else None
    )
    if extension_kind and mime_kind and extension_kind != mime_kind:
        raise ValueError("The file extension does not match its content type.")
    if extension_kind == "pdf" and mime_kind in {"pdf", None}:
        return "pdf"
    if extension_kind == "image" and mime_kind in {"image", None}:
        return "image"
    if extension_kind == "docx" and mime_kind in {"docx", None}:
        return "docx"
    if extension_kind == "text" and mime_kind in {"text", None}:
        return "text"
    raise ValueError(
        "Unsupported file type. Upload a PDF, DOCX, TXT, or an image (PNG, JPG, WEBP)."
    )


def extract_text(data: bytes, filename: str, content_type: str | None) -> tuple[str, str]:
    """Return (text, kind) for an uploaded file."""
    kind = detect_kind(filename, content_type)
    if kind == "pdf":
        text = extract_pdf_text(data)
        if not text:
            raise ExtractionError(
                "This PDF has no extractable text layer. If it is a scan, upload it "
                "as an image instead so it can be read with vision."
            )
        return text, kind

    if kind == "docx":
        text = extract_docx_text(data)
        if not text:
            raise ExtractionError("This DOCX file contains no readable text.")
        return text, kind

    if kind == "text":
        text = extract_plain_text(data)
        if not text:
            raise ExtractionError("This text file is empty.")
        return text, kind

    mime = (content_type or "").split(";")[0].strip().lower()
    if mime not in IMAGE_MIME_TYPES:
        # Trust the extension when the browser sent something generic.
        suffix = filename.rsplit(".", 1)[-1].lower()
        mime = "image/jpeg" if suffix in {"jpg", "jpeg"} else f"image/{suffix}"
    return extract_image_text(data, mime), kind